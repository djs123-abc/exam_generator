"""
试卷 Word 文档生成器 —— 纯 python-docx + 干净模板
严格按照上传试卷模板排版：信息栏 → 得分表 → 题型区
A3横向双栏 / A4单栏，完全无 w15/w14 污染
"""
import os, re
from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from exam_config_schema import FullExamConfig, InfoField
from exam_engine import ExamPaper, Question, QUESTION_TYPES, DEFAULT_INSTRUCTIONS, TYPE_ORDER

# ── 颜色 ─────────────────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x1F, 0x4E, 0x79)
SECONDARY = RGBColor(0x2E, 0x75, 0xB6)
ACCENT    = RGBColor(0x00, 0x70, 0xC0)
GRAY      = RGBColor(0x59, 0x59, 0x59)
LGRAY     = RGBColor(0xAA, 0xAA, 0xAA)
BLACK     = RGBColor(0x00, 0x00, 0x00)

FONT_SONG = '仿宋'
FONT_HEI  = '黑体'
FONT_BSONG = '宋体'

TEMPLATE  = Path(__file__).parent / "clean_template.docx"
ROMAN     = ['一','二','三','四','五','六','七','八','九','十']

# ══════════════════════════════════════════
#  XML 工具函数
# ══════════════════════════════════════════
def _rfonts(run, cn, en=None):
    f = run._r.get_or_add_rPr().get_or_add_rFonts()
    f.set(qn('w:eastAsia'), cn)
    f.set(qn('w:ascii'),    en or cn)
    f.set(qn('w:hAnsi'),    en or cn)
    f.set(qn('w:cs'),       en or cn)

def _spacing(para, before=3, after=3, line=320):
    sp = para._p.get_or_add_pPr().find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        para._p.get_or_add_pPr().append(sp)
    sp.set(qn('w:before'), str(int(before*20)))
    sp.set(qn('w:after'),  str(int(after*20)))
    sp.set(qn('w:lineRule'), 'auto')
    sp.set(qn('w:line'), str(line))

def _indent(para, left=0, first=0):
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); pPr.append(ind)
    if left:  ind.set(qn('w:left'),      str(left))
    if first: ind.set(qn('w:firstLine'), str(first))

def _border_bottom(para, color='2E75B6', sz=10):
    pPr = para._p.get_or_add_pPr()
    pb = pPr.find(qn('w:pBdr'))
    if pb is None: pb = OxmlElement('w:pBdr'); pPr.append(pb)
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '1'); b.set(qn('w:color'), color)
    pb.append(b)

def _border_bottom_gray(para):
    _border_bottom(para, 'CCCCCC', 4)

def _tab_stop(para, pos, align='right'):
    pPr = para._p.get_or_add_pPr()
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None: tabs = OxmlElement('w:tabs'); pPr.append(tabs)
    t = OxmlElement('w:tab')
    t.set(qn('w:val'), align); t.set(qn('w:pos'), str(pos))
    tabs.append(t)

def _run_tab(para):
    r = para.add_run(); r._r.append(OxmlElement('w:tab')); return r

def _page_num(para):
    r = para.add_run()
    for tag, txt in [('w:fldChar','begin'), ('w:instrText',' PAGE '), ('w:fldChar','end')]:
        el = OxmlElement(tag)
        if tag == 'w:instrText': el.text = txt
        else: el.set(qn('w:fldCharType'), txt)
        r._r.append(el)
    _rfonts(r, FONT_SONG); r.font.size = Pt(10)

def _no_border_cell(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{s}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        tcB.append(el)
    tcPr.append(tcB)

def _bottom_border_cell(cell, color='9DC3E6', sz=6):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for s in ('top','left','right','insideH','insideV'):
        el = OxmlElement(f'w:{s}'); el.set(qn('w:val'),'none'); tcB.append(el)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),str(sz))
    bot.set(qn('w:color'),color); tcB.append(bot)
    tcPr.append(tcB)

def _no_border_table(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
    tblB = OxmlElement('w:tblBorders')
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{s}'); el.set(qn('w:val'),'none'); tblB.append(el)
    tblPr.append(tblB)

def _set_cell_width(cell, mm):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    w = tcPr.find(qn('w:tcW'))
    if w is None: w = OxmlElement('w:tcW'); tcPr.append(w)
    w.set(qn('w:w'), str(int(mm * 56.69)))
    w.set(qn('w:type'), 'dxa')

# ══════════════════════════════════════════
#  run/para 快速创建
# ══════════════════════════════════════════
def _run(para, text, bold=False, size=11, color=None, font=None, italic=False):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    _rfonts(r, font or (FONT_HEI if bold else FONT_SONG))
    return r

def _para(doc, text='', bold=False, size=11, color=None, font=None,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          before=3, after=3, line=320, indent_left=0):
    p = doc.add_paragraph()
    p.alignment = align
    _spacing(p, before, after, line)
    if indent_left: _indent(p, left=indent_left)
    if text: _run(p, text, bold=bold, size=size, color=color, font=font)
    return p

def _empty(doc, before=4):
    p = doc.add_paragraph(); _spacing(p, before, 0); return p

# ══════════════════════════════════════════
#  页面设置
# ══════════════════════════════════════════
def _setup_page(doc, config: FullExamConfig):
    section = doc.sections[0]
    if config.paper_size == 'A4':
        section.page_width    = Mm(210)
        section.page_height   = Mm(297)
        section.orientation   = WD_ORIENT.PORTRAIT
    else:  # A3 横向
        section.page_width    = Mm(420)
        section.page_height   = Mm(297)
        section.orientation   = WD_ORIENT.LANDSCAPE
    section.top_margin    = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin   = Mm(15)
    section.right_margin  = Mm(15)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)

    # A3 横向双栏
    if config.paper_size != 'A4' and config.layout == 'double_column':
        sectPr = section._sectPr
        for old in sectPr.findall(qn('w:cols')):
            sectPr.remove(old)
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'),        '2')
        cols.set(qn('w:space'),      '567')
        cols.set(qn('w:equalWidth'), '1')
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is not None: pgMar.addnext(cols)
        else: sectPr.append(cols)

    return section

def _total_w_twp(config: FullExamConfig) -> int:
    if config.paper_size == 'A4':
        return int((210 - 15 - 15) * 56.69)   # 180mm
    return int((420 - 15 - 15) * 56.69)        # 390mm

# ══════════════════════════════════════════
#  页眉 / 页脚
# ══════════════════════════════════════════
def _build_header(section, paper: ExamPaper, total: float, is_ans: bool):
    hdr = section.header
    for p in hdr.paragraphs: p._p.getparent().remove(p._p)
    tw = _total_w_twp(paper.config)
    suffix = '【参考答案与解析】' if is_ans else ''
    p = hdr.add_paragraph()
    _spacing(p, 0, 4); _border_bottom(p, '2E75B6', 10)
    _tab_stop(p, tw, 'right')
    _run(p, f"{paper.title}{suffix}", bold=True, size=13, color=PRIMARY)
    _run_tab(p)
    _run(p, f"科目：{paper.config.subject or '综合'}    时间：{paper.config.exam_time}分钟    总分：{total:.0f}分",
         size=9, color=GRAY)

def _build_footer(section, is_ans: bool, config: FullExamConfig):
    ftr = section.footer
    for p in ftr.paragraphs: p._p.getparent().remove(p._p)
    tw = _total_w_twp(config)
    p = ftr.add_paragraph()
    _spacing(p, 4, 0)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr'); pPr.append(pb)
    top = OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'6')
    top.set(qn('w:color'),'9DC3E6'); pb.append(top)
    _tab_stop(p, tw//2, 'center'); _tab_stop(p, tw, 'right')
    hint = '答案仅供参考' if is_ans else '请在规定时间内完成作答'
    _run(p, hint, size=9, color=LGRAY)
    _run_tab(p)
    _run(p, '第 ', size=10)
    _page_num(p)
    _run(p, ' 页', size=10)
    _run_tab(p)
    _run(p, '答案卷' if is_ans else '试卷', size=9, color=LGRAY)

# ══════════════════════════════════════════
#  信息栏（严格按模板：姓名+单位及职务+所驻镇村）
# ══════════════════════════════════════════
def _build_info_bar(doc, paper: ExamPaper, is_ans: bool):
    cfg = paper.config
    tw_mm = 180 if cfg.paper_size == 'A4' else 390

    if is_ans:
        # 答案卷只显示标题
        p = _para(doc, paper.title, bold=True, size=16, color=PRIMARY,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=4)
        _rfonts(p.runs[0], FONT_HEI)
        p2 = _para(doc, '参考答案与解析', bold=True, size=13, color=ACCENT,
                   align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=8)
        _rfonts(p2.runs[0], FONT_HEI)
        return

    # ── 信息行（姓名 单位/学校 班级 等自定义字段）──────────────────────────
    fields = cfg.info_fields
    total_weight = sum(f.width for f in fields)

    # 先输出信息行文字（仿宋，加粗，仿模板）
    info_para = doc.add_paragraph()
    _spacing(info_para, 0, 4, 300)
    for i, field_cfg in enumerate(fields):
        col_mm = tw_mm * field_cfg.width / total_weight
        blank  = '　' * field_cfg.blank_len
        label_text = f"{field_cfg.label}：{blank}"
        r = _run(info_para, label_text, bold=True, size=11)
        _rfonts(r, FONT_HEI)
        if i < len(fields) - 1:
            _run(info_para, '    ', size=11)   # 间隔

    # ── 试卷大标题 ──────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_title, 2, 2, 340)
    r = _run(p_title, paper.title, bold=True, size=16, color=PRIMARY)
    _rfonts(r, FONT_HEI)

    # ── 得分表格（一/二/三/四/总分）────────────────────────────────────────
    if cfg.score_table.enabled:
        _build_score_table(doc, paper)

    _empty(doc, before=4)


def _build_score_table(doc, paper: ExamPaper):
    """仿模板得分汇总表：各大题分栏 + 总分"""
    cfg = paper.config

    # 确定有哪些题型大组
    groups = {}
    for q in paper.questions:
        groups.setdefault(q.q_type, []).append(q)
    present = [k for k in TYPE_ORDER if k in groups]

    # 列标题：一、二、三… + 总分
    if cfg.score_table.sections:
        headers = list(cfg.score_table.sections)
    else:
        headers = [ROMAN[i] for i in range(len(present))] + ['总分']

    if '总分' not in headers:
        headers.append('总分')

    n = len(headers)
    tw_mm = 180 if cfg.paper_size == 'A4' else 390
    col_mm = tw_mm / n

    tbl = doc.add_table(rows=2, cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_border_table(tbl)

    # 设置外边框
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    tblB = tblPr.find(qn('w:tblBorders'))
    for s in ('top','left','bottom','right','insideH','insideV'):
        el = tblB.find(qn(f'w:{s}'))
        if el is None: el = OxmlElement(f'w:{s}'); tblB.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'6')
        el.set(qn('w:color'),'2E75B6')

    # 行0：标题行
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_width(cell, col_mm)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(cp, 2, 2, 280)
        _run(cp, h, bold=True, size=11, color=PRIMARY)

    # 行1：空白得分行
    for i in range(n):
        cell = tbl.rows[1].cells[i]
        _set_cell_width(cell, col_mm)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(cp, 8, 8, 280)
        _run(cp, '', size=11)

    _empty(doc, before=2)


# ══════════════════════════════════════════
#  题型区块标题
# ══════════════════════════════════════════
def _section_title(doc, idx: int, qt_cfg, questions: list[Question]):
    total = sum(q.score for q in questions)
    per   = questions[0].score if questions else 0
    count = len(questions)
    name  = qt_cfg.custom_name or QUESTION_TYPES.get(qt_cfg.type_key, qt_cfg.type_key)
    inst  = qt_cfg.custom_instruction or DEFAULT_INSTRUCTIONS.get(qt_cfg.type_key, '')

    # 标题行：如"一、单项选择题（每题1.5分，共30分）"
    p = doc.add_paragraph()
    _spacing(p, 8, 3, 320); _border_bottom(p, '2E75B6', 10)
    label = f"{ROMAN[idx]}、{name}（每题{per:.4g}分，共{count}题，共{total:.4g}分）"
    _run(p, label, bold=True, size=12, color=PRIMARY)

    # 说明行
    p2 = doc.add_paragraph()
    _spacing(p2, 2, 4, 300)
    _run(p2, inst, size=10, color=GRAY)


# ══════════════════════════════════════════
#  各题型渲染
# ══════════════════════════════════════════
def _render_choice(doc, q: Question, num: int, show_ans: bool, cfg: FullExamConfig):
    # 题干
    p = doc.add_paragraph()
    _spacing(p, 4, 2, 360)
    bracket = '（    ）' # 选择题括号放题干末尾，仿模板
    score_hint = f'（{q.score:.4g}分）' if cfg.show_score_per_question else ''
    _run(p, f"{num}. ", bold=True, size=11)
    _run(p, q.content)
    _run(p, f"  {bracket}", size=11)
    if score_hint: _run(p, f"  {score_hint}", size=10, color=GRAY)

    # 选项
    opts = q.options or []
    tw_mm = 180 if cfg.paper_size == 'A4' else 190
    if len(opts) <= 2:
        for opt in opts:
            po = doc.add_paragraph(); _spacing(po, 2, 2, 300); _indent(po, left=420)
            _run(po, opt, size=10.5)
    else:
        hw = Mm((tw_mm - 10) / 2)
        for i in range(0, len(opts), 2):
            L, R = opts[i], opts[i+1] if i+1 < len(opts) else ''
            tbl = doc.add_table(rows=1, cols=2)
            _no_border_table(tbl)
            for ci, txt in enumerate([L, R]):
                cell = tbl.rows[0].cells[ci]
                cell.width = hw
                _no_border_cell(cell)
                cp = cell.paragraphs[0]; _spacing(cp, 2, 2, 300); _indent(cp, left=420)
                _run(cp, txt, size=10.5)

    if show_ans:
        pa = doc.add_paragraph(); _spacing(pa, 2, 2, 300); _indent(pa, left=420)
        _run(pa, '【答案】', bold=True, color=ACCENT, size=10.5)
        _run(pa, q.answer, color=ACCENT, size=10.5)
        if q.analysis:
            pe = doc.add_paragraph(); _spacing(pe, 0, 4, 280); _indent(pe, left=420)
            _run(pe, '【解析】', bold=True, color=GRAY, size=10)
            _run(pe, q.analysis, color=GRAY, size=10)


def _render_truefalse(doc, q: Question, num: int, show_ans: bool, cfg: FullExamConfig):
    p = doc.add_paragraph(); _spacing(p, 4, 2, 360)
    score_hint = f'（{q.score:.4g}分）' if cfg.show_score_per_question else ''
    _run(p, f"{num}. ", bold=True, size=11)
    _run(p, q.content)
    _run(p, '  （   ）', size=11)
    if score_hint: _run(p, f"  {score_hint}", size=10, color=GRAY)
    if show_ans:
        pa = doc.add_paragraph(); _spacing(pa, 2, 4, 300); _indent(pa, left=420)
        _run(pa, '【答案】', bold=True, color=ACCENT, size=10.5)
        _run(pa, q.answer, color=ACCENT, size=10.5)
        if q.analysis:
            pe = doc.add_paragraph(); _spacing(pe, 0, 4, 280); _indent(pe, left=420)
            _run(pe, '【解析】', bold=True, color=GRAY, size=10)
            _run(pe, q.analysis, color=GRAY, size=10)


def _render_fillblank(doc, q: Question, num: int, show_ans: bool, cfg: FullExamConfig):
    p = doc.add_paragraph(); _spacing(p, 4, 2, 380)
    score_hint = f'（{q.score:.4g}分）' if cfg.show_score_per_question else ''
    _run(p, f"{num}. ", bold=True, size=11)
    _run(p, q.content)
    if score_hint: _run(p, f"  {score_hint}", size=10, color=GRAY)
    if show_ans:
        pa = doc.add_paragraph(); _spacing(pa, 2, 4, 300); _indent(pa, left=420)
        _run(pa, '【答案】', bold=True, color=ACCENT, size=10.5)
        _run(pa, q.answer, color=ACCENT, size=10.5)


def _render_open(doc, q: Question, num: int, show_ans: bool, cfg: FullExamConfig):
    lines = (q.content or '').split('\n')
    score_hint = f'（{q.score:.4g}分）' if cfg.show_score_per_question else ''
    p = doc.add_paragraph(); _spacing(p, 6, 2, 360)
    _run(p, f"{num}. ", bold=True, size=11)
    _run(p, lines[0])
    if score_hint: _run(p, f"  {score_hint}", size=10, color=GRAY)
    for line in lines[1:]:
        if line.strip():
            pl = doc.add_paragraph(); _spacing(pl, 2, 2, 340); _indent(pl, left=420)
            _run(pl, line, size=11)

    if not show_ans:
        n_lines = 14 if q.q_type == 'case_analysis' else (12 if q.q_type == 'essay' else 7)
        for _ in range(n_lines):
            pb = doc.add_paragraph(); _spacing(pb, 2, 2)
            pb._p.get_or_add_pPr()
            pPr = pb._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr'); pPr.append(pBdr)
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4')
            bot.set(qn('w:color'),'CCCCCC'); pBdr.append(bot)
        _empty(doc, before=4)
    else:
        pa = doc.add_paragraph(); _spacing(pa, 4, 2, 300)
        _run(pa, '【参考答案】', bold=True, color=ACCENT, size=10.5)
        for line in (q.answer or '').split('\n'):
            if line.strip():
                pl = doc.add_paragraph(); _spacing(pl, 2, 2, 320); _indent(pl, left=420)
                _run(pl, line, color=ACCENT, size=10.5)
        if q.analysis:
            pk = doc.add_paragraph(); _spacing(pk, 4, 2, 300)
            _run(pk, '【评分要点】', bold=True, color=GRAY, size=10.5)
            pk2 = doc.add_paragraph(); _spacing(pk2, 0, 4, 300); _indent(pk2, left=420)
            _run(pk2, q.analysis, color=GRAY, size=10)
        _empty(doc, before=4)


def _render_q(doc, q: Question, num: int, show_ans: bool, cfg: FullExamConfig):
    qt = q.q_type
    if qt in ('single_choice','multi_choice'):
        _render_choice(doc, q, num, show_ans, cfg)
    elif qt == 'true_false':
        _render_truefalse(doc, q, num, show_ans, cfg)
    elif qt == 'fill_blank':
        _render_fillblank(doc, q, num, show_ans, cfg)
    else:
        _render_open(doc, q, num, show_ans, cfg)


# ══════════════════════════════════════════
#  主构建函数
# ══════════════════════════════════════════
def _build_doc(paper: ExamPaper, show_ans: bool) -> Document:
    doc = Document(str(TEMPLATE))
    for p in doc.paragraphs: p._p.getparent().remove(p._p)

    section = _setup_page(doc, paper.config)
    total = sum(q.score for q in paper.questions)
    _build_header(section, paper, total, show_ans)
    _build_footer(section, show_ans, paper.config)
    _build_info_bar(doc, paper, show_ans)

    # 分组
    groups: dict[str, list[Question]] = {}
    for q in paper.questions:
        groups.setdefault(q.q_type, []).append(q)

    present = [k for k in TYPE_ORDER if k in groups]
    # 找对应 qt_cfg
    qt_cfg_map = {qt.type_key: qt for qt in paper.config.question_types}

    global_num = 1
    for idx, type_key in enumerate(present):
        qs = groups[type_key]
        qt_cfg = qt_cfg_map.get(type_key)
        if qt_cfg is None:
            from exam_config_schema import QuestionTypeConfig as QTC
            qt_cfg = QTC(type_key=type_key, count=len(qs),
                         score_per=qs[0].score if qs else 0, enabled=True)
        _section_title(doc, idx, qt_cfg, qs)
        for q in qs:
            _render_q(doc, q, global_num, show_ans, paper.config)
            global_num += 1

    return doc


# ══════════════════════════════════════════
#  对外接口
# ══════════════════════════════════════════
def generate_docx(papers: list[ExamPaper], output_dir: str,
                  progress_callback=None) -> list[dict]:
    os.makedirs(output_dir, exist_ok=True)
    results = []
    for i, paper in enumerate(papers):
        if progress_callback:
            progress_callback(i, len(papers), f'生成第{paper.paper_index}套Word文档...')
        safe = re.sub(r'[\\/:*?"<>|]', '_', paper.title)
        suffix = f'_第{paper.paper_index}套' if len(papers) > 1 else ''
        exam_path = os.path.join(output_dir, f'{safe}{suffix}_试题.docx')
        ans_path  = os.path.join(output_dir, f'{safe}{suffix}_答案.docx')
        _build_doc(paper, False).save(exam_path)
        _build_doc(paper, True ).save(ans_path)
        results.append({'paper': exam_path, 'answer': ans_path})
    if progress_callback:
        progress_callback(len(papers), len(papers), 'Word文档生成完成')
    return results
